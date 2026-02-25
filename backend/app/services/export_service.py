import logging
import io
import os
from datetime import datetime
from typing import BinaryIO
from bson.objectid import ObjectId
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

logger = logging.getLogger(__name__)


def _safe_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    return str(value)


def _safe_processing_time(value) -> float:
    try:
        return float(value or 0)
    except Exception:
        return 0.0


def _build_report_text(history_item: dict) -> str:
    lines = []
    lines.append(f"TextAnalyzer - {_safe_text(history_item.get('type', 'unknown')).title()} Report")
    lines.append("=" * 60)
    lines.append("")
    lines.append(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"Processing Type: {_safe_text(history_item.get('type', 'unknown')).upper()}")
    lines.append(f"Processing Time: {_safe_processing_time(history_item.get('processing_time_ms')):.2f}ms")

    input_language = _safe_text(history_item.get('input_language'))
    output_language = _safe_text(history_item.get('output_language'))
    if input_language:
        lines.append(f"Source Language: {input_language}")
        lines.append(f"Target Language: {output_language or 'N/A'}")

    lines.append("")
    lines.append("-" * 60)
    lines.append("INPUT TEXT")
    lines.append("-" * 60)
    lines.append(_safe_text(history_item.get('input_text')))
    lines.append("")
    lines.append("-" * 60)
    lines.append("OUTPUT / RESULT")
    lines.append("-" * 60)
    lines.append(_safe_text(history_item.get('output_text')))
    lines.append("")

    metadata = history_item.get('metadata') or {}
    if metadata:
        lines.append("-" * 60)
        lines.append("ADDITIONAL INFORMATION")
        lines.append("-" * 60)
        for key, value in metadata.items():
            lines.append(f"{_safe_text(key).replace('_', ' ').title()}: {_safe_text(value)}")

    return "\n".join(lines)

class ExportService:
    """Service for exporting processing results in various formats"""

    @staticmethod
    def export_to_txt(history_item: dict) -> str:
        """
        Export processing result to plain text

        Args:
            history_item: Processing history item from MongoDB

        Returns:
            Plain text string
        """
        try:
            return _build_report_text(history_item)

        except Exception as e:
            logger.error(f"TXT export error: {e}")
            raise

    @staticmethod
    def export_to_docx(history_item: dict) -> BinaryIO:
        """
        Export processing result to DOCX (Word document)

        Args:
            history_item: Processing history item from MongoDB

        Returns:
            BytesIO object with DOCX file
        """
        try:
            doc = Document()

            # Header with title
            title = doc.add_heading('TextAnalyzer Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            metadata_para = doc.add_paragraph()
            metadata_para.add_run(f"Generated: ").bold = True
            metadata_para.add_run(f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}\n")

            metadata_para.add_run(f"Processing Type: ").bold = True
            metadata_para.add_run(f"{_safe_text(history_item.get('type', 'unknown')).upper()}\n")

            metadata_para.add_run(f"Processing Time: ").bold = True
            metadata_para.add_run(f"{_safe_processing_time(history_item.get('processing_time_ms')):.2f}ms\n")

            if history_item.get('input_language'):
                metadata_para.add_run(f"Languages: ").bold = True
                metadata_para.add_run(
                    f"{_safe_text(history_item.get('input_language'))} → {_safe_text(history_item.get('output_language', 'N/A'))}\n"
                )

            doc.add_paragraph()  # Spacing

            # Input section
            doc.add_heading('Input Text', level=1)
            input_para = doc.add_paragraph(_safe_text(history_item.get('input_text')))
            input_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_paragraph()  # Spacing

            # Output section
            doc.add_heading('Output / Result', level=1)
            output_para = doc.add_paragraph(_safe_text(history_item.get('output_text')))
            output_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Additional info table (if metadata exists)
            if history_item.get('metadata'):
                doc.add_paragraph()  # Spacing
                doc.add_heading('Additional Information', level=2)

                table_data = [['Metric', 'Value']]
                for key, value in history_item['metadata'].items():
                    table_data.append([_safe_text(key).replace('_', ' ').title(), _safe_text(value)])

                table = doc.add_table(rows=len(table_data), cols=2)
                table.style = 'Light Grid Accent 1'

                for i, row_data in enumerate(table_data):
                    table.rows[i].cells[0].text = row_data[0]
                    table.rows[i].cells[1].text = row_data[1]

            # Footer
            doc.add_paragraph()
            footer_para = doc.add_paragraph('—' * 60)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_text = doc.add_paragraph(
                'This document was generated by TextAnalyzer'
            )
            footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Save to BytesIO
            file_bytes = io.BytesIO()
            doc.save(file_bytes)
            file_bytes.seek(0)

            return file_bytes

        except Exception as e:
            logger.error(f"DOCX export error: {e}")
            raise

    @staticmethod
    def export_to_pdf(history_item: dict) -> BinaryIO:
        """
        Export processing result to PDF

        Args:
            history_item: Processing history item from MongoDB

        Returns:
            BytesIO object with PDF file
        """
        try:
            file_bytes = io.BytesIO()

            # Register Unicode font fallback (for CJK and other non-latin scripts)
            unicode_font = 'Helvetica'
            candidate_fonts = [
                'C:/Windows/Fonts/arialuni.ttf',
                'C:/Windows/Fonts/msgothic.ttc',
                'C:/Windows/Fonts/msyh.ttc',
                'C:/Windows/Fonts/NotoSansCJK-Regular.ttc',
            ]

            for idx, font_path in enumerate(candidate_fonts):
                if os.path.exists(font_path):
                    try:
                        font_name = f'TextAnalyzerUnicode{idx}'
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        unicode_font = font_name
                        break
                    except Exception:
                        continue

            # Create PDF document
            doc = SimpleDocTemplate(
                file_bytes,
                pagesize=A4,
                topMargin=0.5*inch,
                bottomMargin=0.5*inch,
                leftMargin=0.75*inch,
                rightMargin=0.75*inch
            )

            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=unicode_font,
                fontSize=24,
                textColor=colors.HexColor('#0A1F5C'),  # Navy Blue
                spaceAfter=12,
                alignment=TA_CENTER
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontName=unicode_font,
                fontSize=14,
                textColor=colors.HexColor('#0A1F5C'),  # Navy Blue
                spaceAfter=6,
                spaceBefore=6
            )

            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontName=unicode_font,
                fontSize=11,
                textColor=colors.HexColor('#1F2937'),
                alignment=TA_JUSTIFY,
                spaceAfter=6
            )

            # Build content
            content = []

            # Title
            content.append(Paragraph("TextAnalyzer Report", title_style))
            content.append(Spacer(1, 12))

            # Metadata table
            metadata_table_data = [
                ['Generated', datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')],
                ['Processing Type', _safe_text(history_item.get('type', 'unknown')).upper()],
                ['Processing Time', f"{_safe_processing_time(history_item.get('processing_time_ms')):.2f}ms"],
            ]

            if history_item.get('input_language'):
                metadata_table_data.append([
                    'Languages',
                    f"{_safe_text(history_item.get('input_language'))} → {_safe_text(history_item.get('output_language', 'N/A'))}"
                ])

            metadata_table = Table(metadata_table_data, colWidths=[2*inch, 4*inch])
            metadata_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F5F7FA')),  # Soft gray
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#1F2937')),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), unicode_font),
                ('FONTNAME', (1, 0), (1, -1), unicode_font),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ]))

            content.append(metadata_table)
            content.append(Spacer(1, 12))

            # Input section
            content.append(Paragraph("Input Text", heading_style))
            input_text = _safe_text(history_item.get('input_text')).replace('\n', '<br/>')
            content.append(Paragraph(input_text, body_style))
            content.append(Spacer(1, 12))

            # Output section
            content.append(Paragraph("Output / Result", heading_style))
            output_text = _safe_text(history_item.get('output_text')).replace('\n', '<br/>')
            content.append(Paragraph(output_text, body_style))

            # Build PDF
            doc.build(content)
            file_bytes.seek(0)

            return file_bytes

        except Exception as e:
            logger.error(f"PDF export error: {e}")
            raise

    @staticmethod
    def get_filename(processing_type: str, format: str) -> str:
        """
        Generate export filename

        Args:
            processing_type: Type of processing (ocr, grammar, etc.)
            format: Export format (pdf, docx, txt)

        Returns:
            Filename string
        """
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        return f"textanalyzer_{processing_type}_{timestamp}.{format.lower()}"
